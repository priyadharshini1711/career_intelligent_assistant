"""Turn an uploaded file into clean text.

The important detail here is that we *keep line structure*. Resumes and job
descriptions carry most of their meaning in layout -- headings, bullet lists,
one role per block. Flattening everything to a single line (the obvious
`re.sub(r"\\s+", " ", text)` move) destroys the only signal we have for
detecting sections, so normalisation collapses horizontal whitespace and runs
of blank lines but never joins lines together.
"""

import io
import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict

import docx
from pypdf import PdfReader

from app.errors import EmptyDocument, FileTooLarge, UnsupportedFileType

# Bullet glyphs that PDF extraction tends to produce; normalised to "- " so
# the chunker can treat every list the same way.
_BULLET_CHARS = "•‣▪●◦⁃∙·"
_BULLET_RE = re.compile(rf"^[\s{_BULLET_CHARS}\-\*–—]+")
# Space, tab, NBSP, and the U+2000-U+200B block (en/em/thin/hair spaces and
# zero-width space) -- all of which PDF extractors emit in place of a space.
_HORIZONTAL_WS = re.compile(r"[ \t  -​]+")
_BLANK_RUN = re.compile(r"\n{3,}")
# "S K I L L S" -- some PDF generators letter-space headings.
_LETTER_SPACED = re.compile(r"^(?:[A-Za-z]\s){2,}[A-Za-z]$")


@dataclass
class ExtractedDocument:
    text: str
    meta: Dict[str, Any] = field(default_factory=dict)


def validate_upload(filename: str, size_bytes: int, allowed: list, max_bytes: int) -> str:
    """Check extension and size before we spend time parsing. Returns the extension."""
    suffix = Path(filename or "").suffix.lower()
    if suffix not in allowed:
        raise UnsupportedFileType(
            f"'{filename}' is not a supported file type.",
            {"filename": filename, "allowed": allowed},
        )
    if size_bytes > max_bytes:
        raise FileTooLarge(
            f"'{filename}' is {size_bytes / 1_048_576:.1f} MB; the limit is "
            f"{max_bytes / 1_048_576:.0f} MB.",
            {"filename": filename, "size_bytes": size_bytes, "max_bytes": max_bytes},
        )
    return suffix


def normalize(text: str) -> str:
    """Collapse horizontal whitespace and blank-line runs, keep line breaks."""
    # NFKC folds ligatures and full-width characters that PDF extraction emits.
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines = []
    for raw in text.split("\n"):
        line = _HORIZONTAL_WS.sub(" ", raw).strip()
        if _LETTER_SPACED.match(line):
            line = line.replace(" ", "")
        lines.append(line)

    text = "\n".join(lines)
    return _BLANK_RUN.sub("\n\n", text).strip()


def normalize_bullet(line: str) -> str:
    """Strip a leading bullet glyph so bullets compare cleanly as skills."""
    return _BULLET_RE.sub("", line).strip()


def is_bullet(line: str) -> bool:
    return bool(line) and bool(_BULLET_RE.match(line)) and line.strip() not in {"-", "*"}


def _extract_pdf(data: bytes) -> ExtractedDocument:
    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        # An empty-password decrypt covers the common "protected but not
        # really" case; anything else we surface honestly rather than
        # returning a blank document.
        try:
            reader.decrypt("")
        except Exception as exc:  # pragma: no cover - depends on the file
            raise EmptyDocument(
                "This PDF is password protected and could not be read.",
                {"reason": str(exc)},
            ) from exc

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # pragma: no cover - malformed page
            pages.append("")

    return ExtractedDocument(
        text="\n\n".join(pages),
        meta={"extractor": "pypdf", "page_count": len(reader.pages)},
    )


def _extract_docx(data: bytes) -> ExtractedDocument:
    document = docx.Document(io.BytesIO(data))

    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            parts.append("")
            continue
        # Word's built-in heading styles are a free, reliable section signal.
        if paragraph.style is not None and (paragraph.style.name or "").startswith("Heading"):
            parts.extend(["", text.upper() if len(text) < 60 else text, ""])
        else:
            parts.append(text)

    # Tables are common in JD templates and invisible to `.paragraphs`.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return ExtractedDocument(
        text="\n".join(parts),
        meta={
            "extractor": "python-docx",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
    )


def _extract_plaintext(data: bytes) -> ExtractedDocument:
    for encoding in ("utf-8", "utf-16", "cp1252"):
        try:
            return ExtractedDocument(
                text=data.decode(encoding), meta={"extractor": f"text/{encoding}"}
            )
        except UnicodeDecodeError:
            continue
    return ExtractedDocument(
        text=data.decode("latin-1", errors="ignore"), meta={"extractor": "text/latin-1"}
    )


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_plaintext,
    ".md": _extract_plaintext,
}


def extract_text(filename: str, data: bytes, suffix: str) -> ExtractedDocument:
    """Extract and normalise text, raising `EmptyDocument` if nothing usable."""
    extractor = _EXTRACTORS.get(suffix)
    if extractor is None:  # pragma: no cover - validate_upload guards this
        raise UnsupportedFileType(f"No extractor registered for '{suffix}'.")

    try:
        extracted = extractor(data)
    except EmptyDocument:
        raise
    except Exception as exc:
        raise EmptyDocument(
            f"Could not read '{filename}'. The file may be corrupt or in an "
            "unexpected format.",
            {"filename": filename, "reason": str(exc)},
        ) from exc

    extracted.text = normalize(extracted.text)

    # A scanned PDF parses fine and yields almost nothing. Say so explicitly
    # instead of letting the user wonder why every answer is vague.
    if len(extracted.text.split()) < 30:
        raise EmptyDocument(
            f"'{filename}' contains little or no extractable text. If it is a "
            "scanned document, OCR is not supported yet -- please upload a "
            "text-based PDF, DOCX, or TXT file.",
            {"filename": filename, "word_count": len(extracted.text.split())},
        )

    extracted.meta["word_count"] = len(extracted.text.split())
    return extracted
