"""Extraction and chunking.

These are the tests that pay for themselves: everything downstream is only as
good as the text and the chunk boundaries, and both are full of quiet edge
cases that produce plausible-looking garbage rather than an exception.
"""

import io

import pytest
from docx import Document as DocxDocument

from app.errors import EmptyDocument, FileTooLarge, UnsupportedFileType
from app.ingestion.chunking import (
    canonical_section,
    chunk_document,
    section_stats,
    split_sections,
)
from app.ingestion.extract import extract_text, is_bullet, normalize, validate_upload
from app.schemas import DocumentKind
from tests.conftest import JOB_BACKEND_TEXT, RESUME_TEXT

ALLOWED = [".pdf", ".docx", ".txt", ".md"]


class TestValidation:
    def test_rejects_unsupported_extension(self):
        with pytest.raises(UnsupportedFileType):
            validate_upload("resume.exe", 100, ALLOWED, 1000)

    def test_rejects_oversized_file(self):
        with pytest.raises(FileTooLarge):
            validate_upload("resume.pdf", 5000, ALLOWED, 1000)

    def test_accepts_valid_upload(self):
        assert validate_upload("resume.PDF", 500, ALLOWED, 1000) == ".pdf"


class TestNormalize:
    def test_preserves_line_structure(self):
        """The whole section-detection scheme depends on newlines surviving."""
        text = "SKILLS\n- Python\n- SQL"
        assert normalize(text).split("\n") == ["SKILLS", "- Python", "- SQL"]

    def test_collapses_horizontal_whitespace_only(self):
        assert normalize("a    b\n\n\n\nc") == "a b\n\nc"

    def test_joins_letter_spaced_headings(self):
        # Some PDF generators render headings as "S K I L L S".
        assert normalize("S K I L L S") == "SKILLS"

    def test_normalises_unicode_spaces(self):
        assert normalize("Python  developer") == "Python developer"


class TestExtract:
    def test_reads_plain_text(self):
        result = extract_text("resume.txt", RESUME_TEXT.encode(), ".txt")
        # Asserted against structure, not against the sample's contents. An
        # earlier version checked for the candidate's name and broke the moment
        # someone edited the fixture -- a test that fails when the data changes
        # rather than when the code does is noise.
        assert "SUMMARY" in result.text
        assert "EXPERIENCE" in result.text
        assert result.meta["word_count"] > 100
        assert result.meta["extractor"].startswith("text/")

    def test_reads_docx_including_tables(self):
        document = DocxDocument()
        document.add_heading("Requirements", level=1)
        document.add_paragraph("5+ years of Python engineering experience in production.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Location"
        table.rows[0].cells[1].text = "Remote, India"
        for _ in range(6):
            document.add_paragraph("Additional filler text to clear the minimum word count.")

        buffer = io.BytesIO()
        document.save(buffer)

        result = extract_text("job.docx", buffer.getvalue(), ".docx")
        assert "REQUIREMENTS" in result.text
        # Table content is invisible to `.paragraphs` and has to be walked
        # separately -- JD templates put half the detail in tables.
        assert "Remote, India" in result.text

    def test_rejects_document_with_no_extractable_text(self):
        """A scanned PDF parses fine and yields nothing; say so explicitly."""
        with pytest.raises(EmptyDocument):
            extract_text("scan.txt", b"tiny", ".txt")

    def test_falls_back_across_encodings(self):
        text = ("Resume of Jose Garcia. " * 10).encode("cp1252")
        assert "Jose" in extract_text("r.txt", text, ".txt").text


class TestSectionDetection:
    @pytest.mark.parametrize(
        "heading,expected",
        [
            ("TECHNICAL SKILLS", "Skills"),
            ("Work Experience", "Experience"),
            ("What You'll Do", "Responsibilities"),
            ("REQUIREMENTS", "Requirements"),
            ("Nice to have", "Preferred"),
            ("Professional Summary", "Summary"),
            ("About Us", "Company"),
        ],
    )
    def test_maps_heading_variants_to_canonical_names(self, heading, expected):
        assert canonical_section(heading) == expected

    def test_ignores_prose_that_is_not_a_heading(self):
        assert canonical_section("I have experience building payment systems.") is None

    def test_splits_resume_into_expected_sections(self):
        names = {name for name, _ in split_sections(normalize(RESUME_TEXT))}
        assert {"Summary", "Skills", "Experience", "Education"} <= names

    def test_splits_job_description_into_expected_sections(self):
        names = {name for name, _ in split_sections(normalize(JOB_BACKEND_TEXT))}
        assert {"Requirements", "Responsibilities", "Preferred"} <= names


class TestChunking:
    def test_every_chunk_carries_its_section(self):
        chunks = chunk_document("d1", DocumentKind.JOB, "Job", normalize(JOB_BACKEND_TEXT))
        assert chunks
        assert all(chunk.section for chunk in chunks)
        assert "Requirements" in section_stats(chunks)

    def test_chunk_ids_are_unique_and_ordered(self):
        chunks = chunk_document("d1", DocumentKind.RESUME, "Resume", normalize(RESUME_TEXT))
        ids = [chunk.id for chunk in chunks]
        assert len(ids) == len(set(ids))
        assert [chunk.index for chunk in chunks] == list(range(len(chunks)))

    def test_bullets_are_never_split_across_chunks(self):
        """A half-bullet retrieves as a chunk that means neither thing."""
        text = "REQUIREMENTS\n" + "\n".join(
            f"- Requirement number {i} covering a specific and clearly worded capability."
            for i in range(40)
        )
        chunks = chunk_document("d1", DocumentKind.JOB, "Job", text, target_words=40)
        assert len(chunks) > 1
        for chunk in chunks:
            for line in chunk.text.split("\n"):
                if line.startswith("- Requirement"):
                    assert line.rstrip().endswith("capability.")

    def test_respects_target_size_with_tolerance(self):
        chunks = chunk_document(
            "d1", DocumentKind.JOB, "Job", normalize(JOB_BACKEND_TEXT), target_words=60
        )
        # Overlap plus whole-bullet packing means the budget is a target, not a
        # hard cap; the guard is that nothing runs away.
        assert all(chunk.word_count <= 60 * 2.2 for chunk in chunks)

    def test_splits_a_single_oversized_paragraph(self):
        text = "SUMMARY\n" + " ".join(f"word{i}" for i in range(500))
        chunks = chunk_document("d1", DocumentKind.RESUME, "Resume", text, target_words=50)
        assert len(chunks) > 5

    def test_overlap_carries_context_between_chunks(self):
        text = "EXPERIENCE\n" + "\n".join(
            f"- Delivered project {i} using Python and PostgreSQL for the platform team."
            for i in range(20)
        )
        chunks = chunk_document(
            "d1", DocumentKind.RESUME, "Resume", text, target_words=40, overlap_words=15
        )
        assert len(chunks) > 1
        first_words = set(chunks[0].text.split())
        assert first_words & set(chunks[1].text.split())

    def test_keeps_short_sections(self):
        """Regression: a section under the minimum word count was dropped.

        The sample resume's CERTIFICATIONS block is seven words, and it names
        an AWS certification a job description will ask about. It has to be
        retrievable.
        """
        chunks = chunk_document("d1", DocumentKind.RESUME, "Resume", normalize(RESUME_TEXT))
        sections = section_stats(chunks)
        assert "Certifications" in sections
        combined = " ".join(chunk.text for chunk in chunks)
        assert "AWS Certified Solutions Architect" in combined

    def test_short_trailing_block_merges_into_its_section(self):
        text = "EXPERIENCE\n" + "\n".join(
            f"- Delivered project {i} using Python and PostgreSQL for the platform team."
            for i in range(12)
        ) + "\n- Small tail."
        chunks = chunk_document("d1", DocumentKind.RESUME, "Resume", text, target_words=40)
        assert all(chunk.word_count >= 10 for chunk in chunks)
        assert "Small tail." in " ".join(chunk.text for chunk in chunks)

    def test_drops_contact_details_from_resume(self):
        chunks = chunk_document("d1", DocumentKind.RESUME, "Resume", normalize(RESUME_TEXT))
        combined = " ".join(chunk.text for chunk in chunks)
        # PII that adds nothing to retrieval should not reach the prompt.
        assert "priya.ramesh@example.com" not in combined

    def test_bullet_detection(self):
        assert is_bullet("- Built a thing")
        assert is_bullet("• Built a thing")
        assert not is_bullet("Built a thing")
